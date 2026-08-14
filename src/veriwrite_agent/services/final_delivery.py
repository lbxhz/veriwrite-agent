"""Assemble, audit, confirm, and export the final MVP paper."""

from __future__ import annotations

import json
import re
from html import unescape
from datetime import datetime, timezone
from io import BytesIO
from math import ceil

from pydantic import ValidationError

from veriwrite_agent.llm.base import LLMClient, LLMResponseError
from veriwrite_agent.models.final_delivery import (
    FinalMatterProposal,
    FinalPaperAudit,
    FinalPaperAuditIssue,
    FinalPaperPackage,
    FinalReferenceEntry,
)
from veriwrite_agent.models.writing_quality import (
    ManuscriptQualityFinding,
    ManuscriptQualityReview,
)
from veriwrite_agent.models.writing import BodyDraftPackage
from veriwrite_agent.models.writing_handoff import V04WritingHandoff
from veriwrite_agent.services.requirement_policy import (
    RequirementPolicyCompiler,
    ai_generation_prohibitions,
    source_restriction_reasons,
)
from veriwrite_agent.services.pdf_acquisition import (
    evidence_document_identity_conflicts,
)
from veriwrite_agent.services.topic_admission import audit_topic_admission
from veriwrite_agent.services.writing_quality import (
    content_similarity,
    false_self_attribution_detail,
    language_mismatch_detail,
)


class FinalDeliveryError(ValueError):
    """Raised when final delivery would violate a confirmed contract."""


class LLMFinalMatterWriter:
    """Generate title, abstract, keywords, and conclusion from the confirmed body."""

    def __init__(self, client: LLMClient) -> None:
        self._client = client

    def _complete_with_retry(self, messages: list[dict[str, object]]) -> str:
        """Call the provider, retrying transient empty/truncated output once.

        A single truncated response must not abort final-matter generation. The
        retry asks the model to compress rather than truncate; if the provider
        still fails, the resulting ``FinalDeliveryError`` joins the same contract
        error channel the callers already handle.
        """
        last_error: LLMResponseError | None = None
        for attempt in range(2):
            try:
                return self._client.complete(
                    messages,
                    response_format={"type": "json_object"},
                )
            except LLMResponseError as exc:
                last_error = exc
                if attempt == 0:
                    messages = [
                        *messages,
                        {
                            "role": "user",
                            "content": (
                                "The provider returned empty or truncated output. Return "
                                "the complete JSON object, compressing rather than "
                                "truncating so it fits within the configured output limit."
                            ),
                        },
                    ]
        raise FinalDeliveryError(
            "LLM final-matter generation failed after a transient provider retry: "
            f"{last_error}"
        ) from last_error

    def draft(
        self,
        handoff: V04WritingHandoff,
        body: BodyDraftPackage,
    ) -> FinalMatterProposal:
        policy = handoff.requirement_policy or RequirementPolicyCompiler().compile(
            handoff.requirement
        )
        prohibitions = ai_generation_prohibitions(policy)
        if prohibitions:
            raise FinalDeliveryError(
                "AI final-matter generation is prohibited: " + "; ".join(prohibitions)
            )
        schema = json.dumps(
            FinalMatterProposal.model_json_schema(),
            ensure_ascii=False,
            separators=(",", ":"),
        )
        required_fields = _required_final_matter_fields(
            policy.structure.required_sections,
            body.markdown,
        )
        messages = [
            {
                "role": "system",
                "content": (
                    "Return JSON only. Using only the confirmed body, write a final "
                    "paper title, abstract, 3-8 keywords, and conclusion. When required, "
                    "also populate introduction, current_status_analysis, problems, and "
                    "technology_trends from the confirmed body. Each structural field must "
                    "perform a distinct editorial role rather than copy a body paragraph. "
                    "For Chinese papers, keep the abstract around 300-400 Chinese counted "
                    "units and include only scope, reviewed objects, core synthesis, and "
                    "major challenges; omit named instruments and case-level details. The "
                    "introduction must explain the problem, significance, scope, and paper "
                    "structure; leave Aeolus, Himawari-8, GIIRS, datasets, algorithms, and "
                    "specific performance details to the body. Current-status analysis "
                    "must be 2-3 real paragraphs answering: what capability improved, what "
                    "remains unresolved, and what domestic/international differences or "
                    "next directions follow. It must not repeat the data-acquisition survey "
                    "and must not include uncited exact metrics. This is a literature review: "
                    "never claim that 本文、本研究 or 我们 proposed, used, measured, or "
                    "validated methods/results belonging to cited authors. Synthesize the "
                    "body rather than introduce new evidence. Do not add "
                    "citations, DOI values, papers, methods, results, numbers, or claims "
                    "that are absent from the body. "
                    f"The required optional fields for this paper are: {required_fields}. "
                    f"Use the requested output language: {policy.output_language}. "
                    "Required structural sections: "
                    f"{json.dumps(policy.structure.required_sections, ensure_ascii=False)}. "
                    f"Satisfy this schema: {schema}"
                ),
            },
            {"role": "user", "content": body.markdown},
        ]
        raw = self._complete_with_retry(messages)
        proposal: FinalMatterProposal | None = None
        first_error: ValidationError | FinalDeliveryError | None = None
        try:
            proposal = _parse_final_matter(raw)
            _validate_required_final_matter_fields(proposal, required_fields)
            _validate_final_matter_has_no_self_authored_citations(proposal)
        except (ValidationError, FinalDeliveryError) as exc:
            first_error = exc
        if proposal is None:
            repaired_raw = self._complete_with_retry(
                [
                    *messages,
                    {"role": "assistant", "content": raw},
                    {
                        "role": "user",
                        "content": (
                            "Repair the JSON so it satisfies the schema and contains every "
                            f"required optional field: {required_fields}. Preserve all "
                            "supported content and return JSON only. Error: "
                            f"{str(first_error)[:1600]}"
                        ),
                    },
                ],
            )
            try:
                proposal = _parse_final_matter(repaired_raw)
                _validate_required_final_matter_fields(proposal, required_fields)
                _validate_final_matter_has_no_self_authored_citations(proposal)
            except (ValidationError, FinalDeliveryError) as exc:
                raise FinalDeliveryError(
                    "LLM final matter violates the data contract after one repair: "
                    + _final_matter_error_detail(exc)
                ) from exc
        return self._global_edit(
            proposal,
            body,
            required_fields=required_fields,
            output_language=policy.output_language,
            schema=schema,
        )

    def _global_edit(
        self,
        proposal: FinalMatterProposal,
        body: BodyDraftPackage,
        *,
        required_fields: list[str],
        output_language: str,
        schema: str,
    ) -> FinalMatterProposal:
        """Run a separate full-manuscript editorial pass before final assembly."""

        payload = {
            "paper_genre": "literature_review",
            "output_language": output_language,
            "candidate_final_matter": proposal.model_dump(mode="json"),
            "confirmed_body": body.markdown,
        }
        messages = [
            {
                "role": "system",
                "content": (
                    "Act as an independent whole-manuscript editor, not the original writer. "
                    "Return a revised final-matter JSON object only; do not rewrite or quote "
                    "the body and do not create citations. Compare the candidate abstract, "
                    "introduction, current-status analysis, problems, trends, and conclusion "
                    "against every body section. Remove copied or near-duplicate passages and "
                    "make each field perform its own role. The abstract contains only scope, "
                    "review objects, core judgment, and major challenges. The introduction "
                    "contains the research problem, significance, scope, and roadmap, without "
                    "case-level instrument or algorithm details. Current-status analysis must "
                    "be 2-3 concise synthesis paragraphs about capability gains, unresolved "
                    "problems, and domestic/international differences or next directions; it "
                    "must not retell the first technical chapter or contain exact metrics. "
                    "Never attribute a cited study's method or result to 本文、本研究、本论文 "
                    "or 我们. Preserve only claims already supported by the body. Required "
                    f"fields: {required_fields}. Return JSON satisfying: {schema}"
                ),
            },
            {"role": "user", "content": json.dumps(payload, ensure_ascii=False)},
        ]
        last_error: ValidationError | FinalDeliveryError | None = None
        last_edited: FinalMatterProposal | None = None
        for attempt in range(3):
            raw = self._complete_with_retry(messages)
            try:
                edited = _parse_final_matter(raw)
                last_edited = edited
                _validate_required_final_matter_fields(edited, required_fields)
                _validate_final_matter_has_no_self_authored_citations(edited)
                _validate_final_matter_editorial_quality(
                    edited,
                    body.markdown,
                    output_language=output_language,
                )
                return edited
            except (ValidationError, FinalDeliveryError) as exc:
                last_error = exc
                if attempt < 2:
                    messages.extend(
                        [
                            {"role": "assistant", "content": raw},
                            {
                                "role": "user",
                                "content": (
                                    "Revise the final-matter JSON again. Fix every listed "
                                    "structural problem instead of paraphrasing the same "
                                    "content. Return JSON only. Deterministic review: "
                                    f"{_final_matter_error_detail(exc)[:1800]}"
                                ),
                            },
                        ]
                    )
        if last_edited is not None and last_error is not None:
            error_detail = _final_matter_error_detail(last_error)
            repaired = last_edited
            if "abstract" in error_detail or "introduction" in error_detail:
                repaired = self._repair_abstract_and_introduction(
                    repaired,
                    body,
                    output_language=output_language,
                )
            if "current_status_analysis" in error_detail:
                repaired = self._repair_current_status_analysis(
                    repaired,
                    body,
                    output_language=output_language,
                )
            _validate_required_final_matter_fields(repaired, required_fields)
            _validate_final_matter_has_no_self_authored_citations(repaired)
            _validate_final_matter_editorial_quality(
                repaired,
                body.markdown,
                output_language=output_language,
            )
            return repaired
        raise FinalDeliveryError(
            "full-manuscript editorial pass failed after adaptive repair: "
            + _final_matter_error_detail(last_error or FinalDeliveryError("unknown error"))
        )

    def _repair_abstract_and_introduction(
        self,
        proposal: FinalMatterProposal,
        body: BodyDraftPackage,
        *,
        output_language: str,
    ) -> FinalMatterProposal:
        """Give the two front-matter roles a small, independently validated task."""

        payload = {
            "output_language": output_language,
            "abstract_candidate": proposal.abstract,
            "introduction_candidate": proposal.introduction,
            "confirmed_body": body.markdown,
        }
        messages = [
            {
                "role": "system",
                "content": (
                    "Edit only the abstract and introduction of a literature review. Return "
                    "JSON with abstract (one string) and introduction_paragraphs (an array "
                    "of exactly two strings). For a Chinese paper, the abstract must contain "
                    "300-400 counted units and only the research scope, reviewed object "
                    "categories, central synthesis, and major challenges. It must contain no "
                    "named instrument, satellite, dataset, algorithm, paper, author, exact "
                    "metric, year, or case detail. Introduction paragraph 1 explains the "
                    "problem, significance, and scope; paragraph 2 explains the paper's "
                    "argument order. The introduction must also contain no named case, "
                    "algorithm, dataset, metric, or result. Do not copy any body paragraph, "
                    "add citations, or invent evidence. Return JSON only."
                ),
            },
            {"role": "user", "content": json.dumps(payload, ensure_ascii=False)},
        ]
        last_error: Exception | None = None
        last_candidate: FinalMatterProposal | None = None
        for attempt in range(3):
            raw = self._complete_with_retry(messages)
            try:
                parsed = json.loads(raw)
                abstract = " ".join(str(parsed.get("abstract", "")).split())
                paragraphs = parsed.get("introduction_paragraphs")
                if not isinstance(paragraphs, list) or len(paragraphs) != 2:
                    raise FinalDeliveryError(
                        "introduction repair must return exactly two paragraphs"
                    )
                introduction_parts = [
                    " ".join(str(paragraph).split())
                    for paragraph in paragraphs
                    if str(paragraph).strip()
                ]
                if len(introduction_parts) != 2:
                    raise FinalDeliveryError(
                        "introduction repair returned an empty paragraph"
                    )
                repaired = proposal.model_copy(
                    update={
                        "abstract": abstract,
                        "introduction": "\n\n".join(introduction_parts),
                    }
                )
                last_candidate = repaired
                _validate_front_matter_roles(
                    repaired,
                    body.markdown,
                    output_language=output_language,
                )
                return repaired
            except (json.JSONDecodeError, ValidationError, FinalDeliveryError) as exc:
                last_error = exc
                if attempt < 2:
                    messages.extend(
                        [
                            {"role": "assistant", "content": raw},
                            {
                                "role": "user",
                                "content": (
                                    "Repair only these two fields. Remove all named cases, "
                                    "Latin model/instrument names and numbers; keep exactly "
                                    "two introduction paragraphs. Error: "
                                    f"{_final_matter_error_detail(exc)[:1200]}"
                                ),
                            },
                        ]
                    )
        if last_candidate is not None:
            return self._repair_abstract_only(
                last_candidate,
                body,
                output_language=output_language,
            )
        raise FinalDeliveryError(
            "targeted front-matter repair failed: "
            + _final_matter_error_detail(last_error or FinalDeliveryError("unknown error"))
        )

    def _repair_abstract_only(
        self,
        proposal: FinalMatterProposal,
        body: BodyDraftPackage,
        *,
        output_language: str,
    ) -> FinalMatterProposal:
        """Expand or condense an otherwise clean abstract without touching other fields."""

        current_units = _final_text_units(proposal.abstract, output_language)
        messages = [
            {
                "role": "system",
                "content": (
                    "Edit only the abstract of this literature review. Return JSON with "
                    "exactly one key: abstract. For Chinese output, write 340-400 Chinese "
                    "Han characters (not an estimate and not 340 tokens). Cover research "
                    "scope, reviewed object categories, central cross-study judgments, and "
                    "major challenges. Use no named instrument, satellite, dataset, model, "
                    "algorithm, author, paper, exact metric, year, digit, citation, or Latin "
                    "identifier. Do not copy a body paragraph or add evidence. Return JSON "
                    "only."
                ),
            },
            {
                "role": "user",
                "content": json.dumps(
                    {
                        "output_language": output_language,
                        "current_counted_units": current_units,
                        "current_abstract": proposal.abstract,
                        "confirmed_body": body.markdown,
                    },
                    ensure_ascii=False,
                ),
            },
        ]
        last_error: Exception | None = None
        for attempt in range(3):
            raw = self._complete_with_retry(messages)
            abstract = ""
            try:
                parsed = json.loads(raw)
                abstract = " ".join(str(parsed.get("abstract", "")).split())
                repaired = proposal.model_copy(update={"abstract": abstract})
                _validate_front_matter_roles(
                    repaired,
                    body.markdown,
                    output_language=output_language,
                )
                return repaired
            except (json.JSONDecodeError, ValidationError, FinalDeliveryError) as exc:
                last_error = exc
                if attempt < 2:
                    actual = _final_text_units(
                        abstract,
                        output_language,
                    )
                    messages.extend(
                        [
                            {"role": "assistant", "content": raw},
                            {
                                "role": "user",
                                "content": (
                                    f"The deterministic count is {actual}. Rewrite the "
                                    "abstract to 340-400 Chinese Han characters, while "
                                    "keeping all case names, Latin identifiers and digits "
                                    "out. Return only the one-key JSON object. Error: "
                                    f"{_final_matter_error_detail(exc)[:900]}"
                                ),
                            },
                        ]
                    )
        raise FinalDeliveryError(
            "targeted abstract length repair failed: "
            + _final_matter_error_detail(last_error or FinalDeliveryError("unknown error"))
        )

    def _repair_current_status_analysis(
        self,
        proposal: FinalMatterProposal,
        body: BodyDraftPackage,
        *,
        output_language: str,
    ) -> FinalMatterProposal:
        """Repair one structural field without asking the model to rewrite everything."""

        payload = {
            "output_language": output_language,
            "current_status_candidate": proposal.current_status_analysis,
            "confirmed_body": body.markdown,
        }
        messages = [
            {
                "role": "system",
                "content": (
                    "Act only as the current-status-analysis editor for a literature "
                    "review. Return JSON with exactly one key, paragraphs, whose value is "
                    "an array of exactly three non-empty prose strings. Paragraph 1 "
                    "synthesizes where technical capability has improved. Paragraph 2 "
                    "synthesizes unresolved limitations. Paragraph 3 compares domestic and "
                    "international emphases only when supported by the confirmed body, then "
                    "states bounded next directions; if the body cannot support a regional "
                    "comparison, say that comparative evidence is insufficient instead of "
                    "inventing one. Do not retell individual studies, named instruments, "
                    "datasets, or algorithms. Do not include exact performance metrics, "
                    "citations, DOI values, or claims absent from the body. Each paragraph "
                    "must stay below 700 counted units. This is a review: never claim that "
                    "本文、本研究、本论文 or 我们 proposed, used, measured, or validated a "
                    "cited method or result. Return JSON only."
                ),
            },
            {"role": "user", "content": json.dumps(payload, ensure_ascii=False)},
        ]
        last_error: Exception | None = None
        for attempt in range(3):
            raw = self._complete_with_retry(messages)
            try:
                parsed = json.loads(raw)
                paragraphs = parsed.get("paragraphs")
                if not isinstance(paragraphs, list) or len(paragraphs) != 3:
                    raise FinalDeliveryError(
                        "current-status repair must return exactly three paragraphs"
                    )
                cleaned = [
                    " ".join(str(paragraph).split())
                    for paragraph in paragraphs
                    if str(paragraph).strip()
                ]
                if len(cleaned) != 3:
                    raise FinalDeliveryError(
                        "current-status repair returned an empty paragraph"
                    )
                repaired = proposal.model_copy(
                    update={"current_status_analysis": "\n\n".join(cleaned)}
                )
                _validate_final_matter_editorial_quality(
                    repaired,
                    body.markdown,
                    output_language=output_language,
                )
                return repaired
            except (json.JSONDecodeError, ValidationError, FinalDeliveryError) as exc:
                last_error = exc
                if attempt < 2:
                    messages.extend(
                        [
                            {"role": "assistant", "content": raw},
                            {
                                "role": "user",
                                "content": (
                                    "Repair only this JSON. Return exactly three concise "
                                    "paragraph strings and remove every exact metric. Error: "
                                    f"{_final_matter_error_detail(exc)[:1200]}"
                                ),
                            },
                        ]
                    )
        raise FinalDeliveryError(
            "targeted current-status repair failed: "
            + _final_matter_error_detail(last_error or FinalDeliveryError("unknown error"))
        )


class FinalPaperAssembler:
    """Build citations and bibliography in code, then run the release audit."""

    def assemble(
        self,
        *,
        handoff: V04WritingHandoff,
        body: BodyDraftPackage,
        final_matter: FinalMatterProposal,
        ai_declaration: str | None = None,
        manuscript_review: ManuscriptQualityReview | None = None,
    ) -> FinalPaperPackage:
        policy = handoff.requirement_policy or RequirementPolicyCompiler().compile(
            handoff.requirement
        )
        records = {record.doi: record for record in handoff.evidence_library.records}
        missing_dois = [doi for doi in body.source_dois if doi not in records]
        if missing_dois:
            raise FinalDeliveryError(
                "body citations are missing from the evidence library: " + ", ".join(missing_dois)
            )
        citation_key_to_doi = {binding.citation_key: binding.doi for binding in body.citations}
        ordered_dois = list(dict.fromkeys(body.source_dois))
        references = [
            _reference_entry(
                records[doi],
                index=index,
                citation_key=next(
                    key for key, value in citation_key_to_doi.items() if value == doi
                ),
                bibliography_style=policy.references.bibliography_style,
                numeric=_numeric_citations(policy),
            )
            for index, doi in enumerate(ordered_dois, 1)
        ]
        reference_by_key = {entry.citation_key: entry for entry in references}
        transformed_body, unknown_keys = _render_final_citations(
            body.markdown,
            reference_by_key,
            numeric=_numeric_citations(policy),
        )
        transformed_body = _remove_leading_title(transformed_body)
        markdown = _assemble_markdown(
            final_matter,
            transformed_body,
            references,
            ai_declaration=ai_declaration,
            output_language=policy.output_language,
            required_sections=policy.structure.required_sections,
        )
        audit = _audit_final_paper(
            policy=policy,
            handoff=handoff,
            body=body,
            final_matter=final_matter,
            references=references,
            ai_declaration=ai_declaration,
            unknown_citation_keys=unknown_keys,
            markdown=markdown,
            manuscript_review=manuscript_review,
        )
        return FinalPaperPackage(
            status=("needs_revision" if audit.blocking_count else "ready_for_confirmation"),
            requirement_policy=policy,
            title=final_matter.title,
            abstract=final_matter.abstract,
            keywords=final_matter.keywords,
            introduction=final_matter.introduction,
            current_status_analysis=final_matter.current_status_analysis,
            problems=final_matter.problems,
            technology_trends=final_matter.technology_trends,
            body_markdown=transformed_body,
            conclusion=final_matter.conclusion,
            ai_declaration=ai_declaration,
            references=references,
            markdown=markdown,
            audit=audit,
            manuscript_review=manuscript_review,
        )

    def confirm(
        self,
        package: FinalPaperPackage,
        *,
        confirmed_by: str,
    ) -> FinalPaperPackage:
        if package.status != "ready_for_confirmation":
            raise FinalDeliveryError("final paper still has blocking audit issues")
        review_codes = {
            issue.code
            for issue in package.audit.issues
            if issue.severity == "warning"
            and issue.code
            in {
                "theme_element_requires_user_review",
                "original_analysis_requires_user_review",
                "reference_tool_usage_requires_attestation",
            }
        }
        missing_attestations = review_codes - set(package.user_review_attestations)
        if missing_attestations:
            raise FinalDeliveryError(
                "final paper still needs semantic user review: "
                + ", ".join(sorted(missing_attestations))
            )
        name = confirmed_by.strip()
        if not name:
            raise FinalDeliveryError("confirmed_by cannot be blank")
        return package.model_copy(
            update={
                "status": "confirmed",
                "confirmed_by": name,
                "confirmed_at": datetime.now(timezone.utc),
            }
        )


class FinalPaperDocxExporter:
    """Export a confirmed package using the narrative_proposal base preset."""

    def export(self, package: FinalPaperPackage) -> bytes:
        if package.status != "confirmed":
            raise FinalDeliveryError("DOCX export requires a confirmed final paper")
        try:
            from docx import Document
            from docx.enum.section import WD_SECTION_START
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Inches, Mm, Pt, RGBColor
        except ImportError as exc:
            raise FinalDeliveryError("python-docx is required for DOCX export") from exc

        policy = package.requirement_policy
        document = Document()
        section = document.sections[0]
        _ = WD_SECTION_START
        if policy.formatting.paper_size and "letter" in policy.formatting.paper_size.casefold():
            section.page_width = Inches(8.5)
            section.page_height = Inches(11)
        else:
            section.page_width = Mm(210)
            section.page_height = Mm(297)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

        body_font = policy.formatting.body_font or (
            "SimSun" if policy.output_language in {"Chinese", "bilingual"} else "Times New Roman"
        )
        body_size = _font_size(policy.formatting.body_font_size)
        line_spacing = policy.formatting.line_spacing or 1.333
        normal = document.styles["Normal"]
        normal.font.name = body_font
        normal.font.size = Pt(body_size)
        normal._element.rPr.rFonts.set(qn("w:ascii"), body_font)
        normal._element.rPr.rFonts.set(qn("w:hAnsi"), body_font)
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), body_font)
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(8)
        normal.paragraph_format.line_spacing = line_spacing

        heading_tokens = {
            "Heading 1": (16, "2E74B5", 18, 10),
            "Heading 2": (13, "2E74B5", 12, 6),
            "Heading 3": (12, "1F4D78", 8, 4),
        }
        for name, (size, color, before, after) in heading_tokens.items():
            style = document.styles[name]
            style.font.name = body_font
            style.font.size = Pt(size)
            style.font.color.rgb = RGBColor.from_string(color)
            style._element.rPr.rFonts.set(qn("w:eastAsia"), body_font)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)

        document.core_properties.title = package.title
        document.core_properties.subject = (
            "VeriWrite final paper; base preset narrative_proposal; "
            "named override academic_course_paper"
        )
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(12)
        title_run = title.add_run(package.title)
        _set_run_font(title_run, body_font, 18, bold=True)

        english = policy.output_language == "English"
        _add_section(document, "Abstract" if english else "摘要", package.abstract, level=1)
        keyword_paragraph = document.add_paragraph()
        keyword_run = keyword_paragraph.add_run("Keywords: " if english else "关键词：")
        _set_run_font(keyword_run, body_font, body_size, bold=True)
        keyword_values = keyword_paragraph.add_run(
            (", " if english else "；").join(package.keywords)
        )
        _set_run_font(keyword_values, body_font, body_size)
        required_kinds = {
            kind
            for required in policy.structure.required_sections
            if (kind := _canonical_section_kind(required)) is not None
        }
        if package.introduction:
            _add_section(
                document,
                "Introduction" if english else "引言",
                package.introduction,
                level=1,
            )
        if "research_status" in required_kinds:
            document.add_heading(
                "Research Status" if english else "国内外研究现状",
                level=1,
            )
            _append_markdown(
                document,
                _demote_body_headings(package.body_markdown),
                body_font,
                body_size,
                numeric_superscript=_numeric_citations(policy),
            )
        else:
            _append_markdown(
                document,
                package.body_markdown,
                body_font,
                body_size,
                numeric_superscript=_numeric_citations(policy),
            )
        if package.current_status_analysis:
            _add_section(
                document,
                "Current Status Analysis" if english else "现状分析",
                package.current_status_analysis,
                level=1,
            )
        if package.problems:
            _add_section(
                document,
                "Existing Problems" if english else "存在问题",
                package.problems,
                level=1,
            )
        if package.technology_trends and not _section_present(
            "技术发展趋势", package.body_markdown
        ):
            _add_section(
                document,
                "Technology Trends" if english else "技术发展趋势",
                package.technology_trends,
                level=1,
            )
        conclusion_heading = "Conclusion" if english else "结论"
        if not english and any(
            re.sub(r"\s+", "", required) == "结语"
            for required in policy.structure.required_sections
        ):
            conclusion_heading = "结语"
        _add_section(
            document,
            conclusion_heading,
            package.conclusion,
            level=1,
        )
        document.add_heading("References" if english else "参考文献", level=1)
        for entry in package.references:
            paragraph = document.add_paragraph(_clean_reference_text(entry.formatted_text))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.first_line_indent = Inches(-0.25)
            paragraph.paragraph_format.line_spacing = 1.15
            paragraph.paragraph_format.space_after = Pt(4)
        if package.ai_declaration:
            _add_section(
                document,
                "AI Usage Declaration" if english else "AI 使用声明",
                package.ai_declaration,
                level=1,
            )

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run()
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), "PAGE")
        run._r.append(field)

        output = BytesIO()
        document.save(output)
        return output.getvalue()


def _audit_final_paper(
    *,
    policy,
    handoff: V04WritingHandoff,
    body: BodyDraftPackage,
    final_matter: FinalMatterProposal,
    references: list[FinalReferenceEntry],
    ai_declaration: str | None,
    unknown_citation_keys: list[str],
    markdown: str,
    manuscript_review: ManuscriptQualityReview | None = None,
) -> FinalPaperAudit:
    issues: list[FinalPaperAuditIssue] = []
    identity_conflicts = evidence_document_identity_conflicts(
        handoff.evidence_library
    )
    for expected_doi, detected_dois in identity_conflicts.items():
        issues.append(
            _issue(
                "document_identity_mismatch",
                "blocking",
                "evidence_library.documents",
                (
                    f"文献 {expected_doi} 绑定的 PDF 首页实际检测到 DOI "
                    f"{', '.join(detected_dois)}；该文献的证据卡、正文引文和编号均不可信，"
                    "必须返回 V0.3 重新核验 PDF 身份。"
                ),
            )
        )
    admission = audit_topic_admission(
        handoff.evidence_library,
        policy,
        valid_section_ids=(
            section.section_id for section in handoff.outline.outline.sections
        ),
    )
    if not admission.passed:
        issues.append(
            _issue(
                "topic_admission_incomplete",
                "blocking",
                "literature.topic_admission",
                admission.detail,
            )
        )
    counted = body.counted_words
    body_prose = re.sub(r"\[@[^\]]+\]", "", body.markdown)
    language_detail = language_mismatch_detail(
        body_prose,
        output_language=policy.output_language,
    )
    if language_detail:
        issues.append(
            _issue(
                "body_language_mismatch",
                "blocking",
                "output_language",
                language_detail,
            )
        )
    for finding in manuscript_review.findings if manuscript_review else []:
        issues.append(
            _issue(
                f"body_{finding.code}",
                _manuscript_finding_audit_severity(finding),
                "writing.global_manuscript_review",
                (
                    f"{finding.section_id}:{finding.paragraph_number}="
                    f"{finding.detail} 修订要求：{finding.revision_instruction}"
                ),
            )
        )
    if manuscript_review and manuscript_review.review_status == "deterministic_fallback":
        issues.append(
            _issue(
                "global_manuscript_review_fallback",
                "warning",
                "writing.global_manuscript_review",
                "独立全文审稿模型返回异常；已执行确定性重复、超长段落与错误归属检查。",
            )
        )
    for section_id, paragraph_number, detail in _body_false_attributions(body):
        identity = f"{section_id}:{paragraph_number}="
        if any(
            issue.code == "body_false_self_attribution"
            and issue.detail.startswith(identity)
            for issue in issues
        ):
            continue
        issues.append(
            _issue(
                "body_false_self_attribution",
                "blocking",
                "writing.review_genre",
                identity + detail,
            )
        )
    if policy.length.minimum_units is not None and counted < policy.length.minimum_units:
        issues.append(
            _issue(
                "length_below_minimum",
                "blocking",
                "length",
                f"required >= {policy.length.minimum_units}; actual={counted}",
            )
        )
    if policy.length.maximum_units is not None and counted > policy.length.maximum_units:
        issues.append(
            _issue(
                "length_above_maximum",
                "blocking",
                "length",
                f"required <= {policy.length.maximum_units}; actual={counted}",
            )
        )
    if len(references) < policy.references.minimum_total:
        issues.append(
            _issue(
                "reference_count_below_minimum",
                "blocking",
                "references.minimum_total",
                f"required={policy.references.minimum_total}; actual={len(references)}",
            )
        )
    elif len(references) < policy.references.target_total:
        issues.append(
            _issue(
                "reference_count_below_target",
                (
                    "warning"
                    if policy.references.target_is_approximate
                    else "blocking"
                ),
                "references.target_total",
                (
                    f"target={policy.references.target_total}; "
                    f"actual={len(references)}; "
                    f"origin={policy.references.target_origin}"
                ),
            )
        )
    foreign_count = sum(entry.is_foreign for entry in references)
    required_foreign = (
        ceil(len(references) * policy.references.minimum_foreign_ratio)
        if policy.references.minimum_foreign_ratio is not None
        else None
    )
    if required_foreign is not None and foreign_count < required_foreign:
        issues.append(
            _issue(
                "foreign_ratio_below_minimum",
                "blocking",
                "references.minimum_foreign_ratio",
                f"required_count={required_foreign}; actual={foreign_count}",
            )
        )
    if policy.references.year_from is not None:
        outside = [
            entry.doi
            for entry in references
            if entry.year < policy.references.year_from
            or (policy.references.year_to is not None and entry.year > policy.references.year_to)
        ]
        if outside:
            issues.append(
                _issue(
                    "hard_year_window_violation",
                    "blocking",
                    "references.recent_year_window",
                    ", ".join(outside),
                )
            )
    for entry in references:
        reasons = source_restriction_reasons(
            policy, publisher=entry.publisher, journal=entry.journal, source_type=entry.source_type
        )
        if reasons:
            issues.append(
                _issue(
                    "source_restriction_violation",
                    "blocking",
                    "references.restriction_rules",
                    f"{entry.doi}: {', '.join(reasons)}",
                )
            )
    if unknown_citation_keys:
        issues.append(
            _issue(
                "unknown_citation_key", "blocking", "references", ", ".join(unknown_citation_keys)
            )
        )
    visible_citation_keys = set(
        re.findall(r"@([a-z0-9][a-z0-9_]{2,79})", body.markdown)
    )
    uncited_entries = [
        entry for entry in references if entry.citation_key not in visible_citation_keys
    ]
    if uncited_entries:
        issues.append(
            _issue(
                "uncited_bibliography_item",
                "blocking",
                "references",
                "文后条目未在正文出现："
                + ", ".join(
                    f"[{entry.index}] {entry.doi}" for entry in uncited_entries
                ),
            )
        )
    cluster_limit = policy.references.max_references_per_citation_cluster
    if cluster_limit is not None:
        clusters: dict[tuple[str, int], set[str]] = {}
        for citation in body.citations:
            clusters.setdefault((citation.section_id, citation.paragraph_number), set()).add(
                citation.doi
            )
        excessive = [
            f"{section}:{paragraph}={len(dois)}"
            for (section, paragraph), dois in clusters.items()
            if len(dois) > cluster_limit
        ]
        if excessive:
            issues.append(
                _issue(
                    "citation_cluster_too_large",
                    "blocking",
                    "references.max_references_per_citation_cluster",
                    ", ".join(excessive),
                )
            )
    required_sections = policy.structure.required_sections
    for section in required_sections:
        if not _section_present(section, markdown):
            issues.append(
                _issue(
                    "required_section_missing", "blocking", "structure.required_sections", section
                )
            )
    for deliverable in policy.deliverables:
        normalized_deliverable = deliverable.casefold()
        deliverable_missing = (
            ("正文" in normalized_deliverable or "body" in normalized_deliverable)
            and not body.markdown.strip()
        ) or (
            (
                "参考文献" in normalized_deliverable
                or "reference" in normalized_deliverable
            )
            and not references
        )
        if deliverable_missing:
            issues.append(
                _issue(
                    "required_deliverable_missing",
                    "blocking",
                    "deliverables",
                    deliverable,
                )
            )
    compact_markdown = re.sub(r"\s+", "", markdown.casefold())
    for theme_element in policy.required_theme_elements:
        normalized_element = re.sub(r"\s+", "", theme_element.casefold())
        if normalized_element and normalized_element not in compact_markdown:
            issues.append(
                _issue(
                    "theme_element_requires_user_review",
                    "warning",
                    "required_theme_elements",
                    theme_element,
                )
            )
    if policy.structure.must_include_original_analysis:
        issues.append(
            _issue(
                "original_analysis_requires_user_review",
                "warning",
                "structure.must_include_original_analysis",
                "User must confirm that synthesis contains original analysis.",
            )
        )
    if policy.references.required_management_tools:
        issues.append(
            _issue(
                "reference_tool_usage_requires_attestation",
                "warning",
                "references.required_management_tools",
                ", ".join(policy.references.required_management_tools),
            )
        )
    if policy.ai_usage.declaration_required and not ai_declaration:
        issues.append(
            _issue(
                "ai_declaration_missing",
                "blocking",
                "ai_policy.declaration_required",
                "A confirmed AI usage declaration is required.",
            )
        )
    for unresolved in policy.unresolved_requirements:
        if _resolved_by_executable_policy(policy, unresolved):
            continue
        severity = (
            "blocking"
            if (
                "output_language" in unresolved
                or "bibliography_style" in unresolved
                or unresolved.startswith("hard_rule_not_fully_executable:")
            )
            else "warning"
        )
        issues.append(
            _issue("unresolved_requirement", severity, "unresolved_requirements", unresolved)
        )
    if policy.output_language == "pending_confirmation":
        issues.append(
            _issue(
                "output_language_unconfirmed",
                "blocking",
                "output_language",
                "Final document language is not confirmed.",
            )
        )
    return FinalPaperAudit(
        policy_fingerprint=policy.requirement_fingerprint,
        counted_units=counted,
        reference_count=len(references),
        foreign_reference_count=foreign_count,
        issues=issues,
    )


def _manuscript_finding_audit_severity(
    finding: ManuscriptQualityFinding,
) -> str:
    """Turn high-confidence editor actions into the existing repair workflow."""

    if finding.disposition == "targeted_repair":
        return "blocking"
    return finding.severity


def _resolved_by_executable_policy(policy, unresolved: str) -> bool:
    compact = re.sub(r"\s+", "", unresolved)
    has_minimum_wording = any(token in compact for token in ("至少", "以上"))
    has_approximate_wording = any(token in compact for token in ("左右", "约"))
    return (
        policy.length.minimum_units is not None
        and has_minimum_wording
        and has_approximate_wording
    )


def _issue(code: str, severity: str, path: str, detail: str) -> FinalPaperAuditIssue:
    return FinalPaperAuditIssue(code=code, severity=severity, requirement_path=path, detail=detail)


def _reference_entry(
    record, *, index: int, citation_key: str, bibliography_style: str, numeric: bool
) -> FinalReferenceEntry:
    authors = [_clean_reference_text(author) for author in record.authors] or ["Anonymous"]
    author_text = ", ".join(authors)
    title = _clean_reference_text(record.title)
    journal = _clean_reference_text(record.journal or record.publisher or "Unknown source")
    doi_url = f"https://doi.org/{record.doi}"
    if "gb/t" in bibliography_style.casefold() or "7714" in bibliography_style:
        formatted = f"[{index}] {author_text}. {title}[J]. {journal}, {record.year}. DOI:{record.doi}."
    elif numeric:
        formatted = f"[{index}] {author_text}. {title}. {journal}. {record.year}. {doi_url}."
    else:
        formatted = f"{author_text} ({record.year}). {title}. {journal}. {doi_url}."
    return FinalReferenceEntry(
        citation_key=citation_key,
        index=index,
        doi=record.doi,
        authors=authors,
        year=record.year,
        title=title,
        journal=journal,
        publisher=record.publisher,
        source_type=record.source_type,
        is_foreign=record.is_foreign,
        formatted_text=formatted,
    )


def _clean_reference_text(value: str) -> str:
    """Remove presentation markup returned by metadata providers before export."""

    cleaned = unescape(value)
    cleaned = re.sub(r"<\s*/?\s*(?:sub|sup|i|b|em|strong)\s*>", "", cleaned, flags=re.I)
    return re.sub(r"\s+", " ", cleaned).strip()


def _numeric_citations(policy) -> bool:
    return (
        policy.references.in_text_style == "numeric_superscript"
        or "gb/t" in policy.references.bibliography_style.casefold()
        or "7714" in policy.references.bibliography_style
    )


def _render_final_citations(
    markdown: str, references: dict[str, FinalReferenceEntry], *, numeric: bool
) -> tuple[str, list[str]]:
    unknown: list[str] = []

    def replace(match: re.Match[str]) -> str:
        rendered: list[str] = []
        for raw in match.group(1).split(";"):
            item = raw.strip()
            parsed = re.fullmatch(r"@([a-z0-9_]+)(?:,\s*(.+))?", item)
            if parsed is None:
                return match.group(0)
            key, _locator = parsed.groups()
            reference = references.get(key)
            if reference is None:
                unknown.append(key)
                return match.group(0)
            if numeric:
                rendered.append(str(reference.index))
            else:
                surname = _surname(reference.authors[0] if reference.authors else "Anonymous")
                rendered.append(f"{surname}, {reference.year}")
        return f"[{'; '.join(rendered)}]" if numeric else f"({'; '.join(rendered)})"

    return re.sub(r"\[(@[^\]]+)\]", replace, markdown), list(dict.fromkeys(unknown))


def _surname(author: str) -> str:
    return author.split(",", 1)[0].strip() if "," in author else author.split()[-1]


def _remove_leading_title(markdown: str) -> str:
    return re.sub(r"^#\s+[^\n]+\n+", "", markdown.strip())


def _assemble_markdown(
    final_matter: FinalMatterProposal,
    body: str,
    references: list[FinalReferenceEntry],
    *,
    ai_declaration: str | None,
    output_language: str,
    required_sections: list[str],
) -> str:
    english = output_language == "English"
    abstract_heading = "Abstract" if english else "摘要"
    keyword_label = "Keywords:" if english else "关键词："
    keyword_separator = ", " if english else "；"
    conclusion_heading = "Conclusion" if english else "结论"
    references_heading = "References" if english else "参考文献"
    required_kinds = {
        kind
        for section in required_sections
        if (kind := _canonical_section_kind(section)) is not None
    }
    if not english and any(
        re.sub(r"\s+", "", section) == "结语" for section in required_sections
    ):
        conclusion_heading = "结语"
    parts = [
        f"# {final_matter.title}",
        f"## {abstract_heading}",
        final_matter.abstract,
        f"**{keyword_label}** {keyword_separator.join(final_matter.keywords)}",
    ]
    if final_matter.introduction:
        parts.extend(
            [
                "## Introduction" if english else "## 引言",
                final_matter.introduction,
            ]
        )
    if "research_status" in required_kinds:
        parts.extend(
            [
                "## Research Status" if english else "## 国内外研究现状",
                _demote_body_headings(body),
            ]
        )
    else:
        parts.append(body)
    if final_matter.current_status_analysis:
        parts.extend(
            [
                "## Current Status Analysis" if english else "## 现状分析",
                final_matter.current_status_analysis,
            ]
        )
    if final_matter.problems:
        parts.extend(
            [
                "## Existing Problems" if english else "## 存在问题",
                final_matter.problems,
            ]
        )
    if final_matter.technology_trends and not _section_present(
        "技术发展趋势", body
    ):
        parts.extend(
            [
                "## Technology Trends" if english else "## 技术发展趋势",
                final_matter.technology_trends,
            ]
        )
    parts.extend(
        [
            f"## {conclusion_heading}",
            final_matter.conclusion,
            f"## {references_heading}",
            "\n".join(entry.formatted_text for entry in references),
        ]
    )
    if ai_declaration:
        parts.extend(
            [
                "## AI Usage Declaration" if english else "## AI 使用声明",
                ai_declaration,
            ]
        )
    return "\n\n".join(part.strip() for part in parts if part and part.strip()) + "\n"


def _section_present(required: str, markdown: str) -> bool:
    normalized = re.sub(r"\s+", "", required.casefold())
    canonical = _canonical_section_kind(required)
    if canonical == "problems_and_trends":
        return _section_present("存在问题", markdown) and _section_present(
            "技术发展趋势", markdown
        )
    aliases = {
        "标题": "#",
        "title": "#",
        "摘要": "##摘要",
        "abstract": "##abstract",
        "关键词": "关键词",
        "keywords": "keywords:",
        "正文": "##",
        "body": "##",
        "结论": "##结论",
        "conclusion": "##conclusion",
        "参考文献": "##参考文献",
        "references": "##references",
    }
    compact_markdown = re.sub(r"\s+", "", markdown.casefold())
    if normalized in aliases:
        marker = aliases[normalized]
        if marker in compact_markdown:
            return True
        bilingual_markers = {
            "##摘要": "##abstract",
            "关键词": "keywords:",
            "##结论": "##conclusion",
            "##参考文献": "##references",
        }
        counterpart = bilingual_markers.get(marker)
        if counterpart is None:
            counterpart = next(
                (left for left, right in bilingual_markers.items() if right == marker),
                None,
            )
        return counterpart in compact_markdown if counterpart else False
    headings = [
        re.sub(r"\s+", "", value.casefold())
        for value in re.findall(r"^#{1,3}\s+(.+)$", markdown, re.MULTILINE)
    ]
    canonical_aliases = {
        "introduction": ("引言", "introduction"),
        "research_status": ("研究现状", "researchstatus", "literaturereview"),
        "current_status_analysis": ("现状分析", "currentstatusanalysis"),
        "problems": ("存在问题", "existingproblems", "problems"),
        "trends": ("技术发展趋势", "发展趋势", "technologytrends"),
        "conclusion": ("结语", "结论", "conclusion"),
    }
    if canonical in canonical_aliases:
        return any(
            alias in heading
            for heading in headings
            for alias in canonical_aliases[canonical]
        )
    return any(normalized in heading or heading in normalized for heading in headings)


def _parse_final_matter(raw: str) -> FinalMatterProposal:
    try:
        return FinalMatterProposal.model_validate_json(raw)
    except ValidationError as strict_error:
        try:
            payload = json.loads(raw, strict=False)
        except json.JSONDecodeError:
            raise strict_error
        return FinalMatterProposal.model_validate(payload)


def _required_final_matter_fields(
    required_sections: list[str],
    body_markdown: str,
) -> list[str]:
    kinds = {
        kind
        for section in required_sections
        if (kind := _canonical_section_kind(section)) is not None
    }
    required: list[str] = []
    if "introduction" in kinds and not _section_present("引言", body_markdown):
        required.append("introduction")
    if "current_status_analysis" in kinds and not _section_present(
        "现状分析", body_markdown
    ):
        required.append("current_status_analysis")
    if (
        "problems" in kinds or "problems_and_trends" in kinds
    ) and not _section_present("存在问题", body_markdown):
        required.append("problems")
    if (
        "trends" in kinds or "problems_and_trends" in kinds
    ) and not _section_present("技术发展趋势", body_markdown):
        required.append("technology_trends")
    return required


def _validate_required_final_matter_fields(
    proposal: FinalMatterProposal,
    required_fields: list[str],
) -> None:
    missing = [field for field in required_fields if not getattr(proposal, field)]
    if missing:
        raise FinalDeliveryError(
            "final matter is missing required structural fields: " + ", ".join(missing)
        )


def _validate_final_matter_has_no_self_authored_citations(
    proposal: FinalMatterProposal,
) -> None:
    text_fields = [
        proposal.title,
        proposal.abstract,
        proposal.introduction,
        proposal.current_status_analysis,
        proposal.problems,
        proposal.technology_trends,
        proposal.conclusion,
    ]
    if any(
        value
        and re.search(
            r"\[@|https?://(?:dx\.)?doi\.org/|doi\s*:|10\.\d{4,9}/\S+",
            value,
            re.IGNORECASE,
        )
        for value in text_fields
    ):
        raise FinalDeliveryError(
            "final matter contains a self-authored citation marker or DOI value"
        )


def _validate_final_matter_editorial_quality(
    proposal: FinalMatterProposal,
    body_markdown: str,
    *,
    output_language: str,
) -> None:
    """Enforce distinct section roles after the independent global edit."""

    problems: list[str] = []
    abstract_units = _final_text_units(proposal.abstract, output_language)
    if output_language == "Chinese" and not 250 <= abstract_units <= 450:
        problems.append(
            f"Chinese abstract has {abstract_units} counted units; expected about 300-400 "
            "(accepted range 250-450)"
        )
    elif output_language == "English" and not 120 <= abstract_units <= 260:
        problems.append(
            f"English abstract has {abstract_units} words; accepted range is 120-260"
        )

    _collect_front_matter_role_problems(
        proposal,
        body_markdown,
        output_language=output_language,
        problems=problems,
    )

    named_fields = {
        "abstract": proposal.abstract,
        "introduction": proposal.introduction,
        "current_status_analysis": proposal.current_status_analysis,
        "problems": proposal.problems,
        "technology_trends": proposal.technology_trends,
        "conclusion": proposal.conclusion,
    }
    for field_name, value in named_fields.items():
        if value and (detail := false_self_attribution_detail(value)):
            problems.append(f"{field_name}: {detail}")

    if proposal.current_status_analysis:
        status_paragraphs = _split_prose_paragraphs(proposal.current_status_analysis)
        if not 2 <= len(status_paragraphs) <= 3:
            problems.append(
                "current_status_analysis must contain 2-3 concise synthesis paragraphs; "
                f"actual={len(status_paragraphs)}"
            )
        status_units = [
            _final_text_units(paragraph, output_language)
            for paragraph in status_paragraphs
        ]
        if any(units > 900 for units in status_units) or sum(status_units) > 1800:
            problems.append(
                "current_status_analysis is oversized; each paragraph must be <=900 "
                "counted units and the section <=1800"
            )
        if re.search(r"(?<![A-Za-z])\d+(?:\.\d+)?\s*(?:%|米|公里|km|m\b|nm\b)", proposal.current_status_analysis):
            problems.append(
                "current_status_analysis contains exact uncited performance metrics; "
                "replace them with evidence-bounded synthesis"
            )
        if output_language == "Chinese" and (
            terms := _case_level_terms(proposal.current_status_analysis)
        ):
            problems.append(
                "current_status_analysis contains named case-level terms "
                f"({', '.join(terms[:6])}); synthesize capabilities and limits instead"
            )

    body_paragraphs = _body_prose_paragraphs(body_markdown)
    role_fields = {
        key: value
        for key, value in named_fields.items()
        if key in {"abstract", "introduction", "current_status_analysis"} and value
    }
    for field_name, value in role_fields.items():
        best = max(
            (content_similarity(value, paragraph) for paragraph in body_paragraphs),
            default=0.0,
        )
        if best >= 0.78:
            problems.append(
                f"{field_name} repeats a body paragraph at similarity {best:.0%}; "
                "rewrite it for its distinct structural role"
            )
    role_items = list(role_fields.items())
    for left in range(len(role_items)):
        for right in range(left + 1, len(role_items)):
            left_name, left_text = role_items[left]
            right_name, right_text = role_items[right]
            similarity = content_similarity(left_text, right_text)
            if similarity >= 0.74:
                problems.append(
                    f"{left_name} and {right_name} overlap at {similarity:.0%}; "
                    "separate their rhetorical functions"
                )
    if problems:
        raise FinalDeliveryError("; ".join(problems[:10]))


def _validate_front_matter_roles(
    proposal: FinalMatterProposal,
    body_markdown: str,
    *,
    output_language: str,
) -> None:
    problems: list[str] = []
    abstract_units = _final_text_units(proposal.abstract, output_language)
    if output_language == "Chinese" and not 250 <= abstract_units <= 450:
        problems.append(
            f"Chinese abstract has {abstract_units} counted units; accepted range 250-450"
        )
    elif output_language == "English" and not 120 <= abstract_units <= 260:
        problems.append(
            f"English abstract has {abstract_units} words; accepted range 120-260"
        )
    _collect_front_matter_role_problems(
        proposal,
        body_markdown,
        output_language=output_language,
        problems=problems,
    )
    if problems:
        raise FinalDeliveryError("; ".join(problems[:8]))


def _collect_front_matter_role_problems(
    proposal: FinalMatterProposal,
    body_markdown: str,
    *,
    output_language: str,
    problems: list[str],
) -> None:
    if proposal.introduction and not _has_introduction_roadmap(
        proposal.introduction,
        output_language=output_language,
    ):
        problems.append(
            "introduction is missing an explicit paper-structure roadmap; explain "
            "the order in which the review develops its argument"
        )
    if output_language == "Chinese":
        for field_name, value in (
            ("abstract", proposal.abstract),
            ("introduction", proposal.introduction),
        ):
            terms = _case_level_terms(value or "")
            if terms:
                problems.append(
                    f"{field_name} contains named case-level terms "
                    f"({', '.join(terms[:6])}); keep only its structural role"
                )
            if value and re.search(r"\d", value):
                problems.append(
                    f"{field_name} contains case-level numbers; move them to the body"
                )
    body_paragraphs = _body_prose_paragraphs(body_markdown)
    for field_name, value in (
        ("abstract", proposal.abstract),
        ("introduction", proposal.introduction),
    ):
        if not value:
            continue
        best = max(
            (content_similarity(value, paragraph) for paragraph in body_paragraphs),
            default=0.0,
        )
        if best >= 0.78:
            problems.append(
                f"{field_name} repeats a body paragraph at similarity {best:.0%}; "
                "rewrite it for its distinct structural role"
            )


def _has_introduction_roadmap(text: str, *, output_language: str) -> bool:
    """Return whether an introduction explicitly previews the paper's argument order."""

    compact = " ".join(text.split())
    if output_language == "Chinese":
        patterns = (
            r"(?:本文|本综述|全文|文章).{0,48}(?:结构|安排|分为|组织|依次)",
            r"(?:下文|以下).{0,36}(?:依次|分为|讨论|综述|介绍|展开)",
            r"(?:首先|第一).{0,220}(?:其次|随后|接着).{0,220}(?:最后|最终)",
        )
    else:
        patterns = (
            r"(?:this paper|this review|the paper|the review).{0,80}"
            r"(?:is organized|is structured|is divided|proceeds)",
            r"(?:first|firstly).{0,260}(?:next|then|second).{0,260}"
            r"(?:finally|lastly)",
        )
    return any(re.search(pattern, compact, flags=re.IGNORECASE) for pattern in patterns)


def _case_level_terms(text: str) -> list[str]:
    """Return Latin instrument/model identifiers embedded in Chinese synthesis prose."""

    return list(
        dict.fromkeys(
            token
            for token in re.findall(
                r"(?<![A-Za-z])[A-Za-z][A-Za-z0-9]*(?:[-–][A-Za-z0-9]+)*(?![A-Za-z])",
                text,
            )
            if len(token) >= 4
        )
    )


def _split_prose_paragraphs(text: str) -> list[str]:
    return [
        paragraph.strip()
        for paragraph in re.split(r"\n\s*\n|\r\n\s*\r\n", text)
        if paragraph.strip()
    ]


def _body_prose_paragraphs(markdown: str) -> list[str]:
    return [
        re.sub(r"\s*\[@[^\]]+\]\s*$", "", block.strip())
        for block in re.split(r"\n\s*\n", markdown)
        if block.strip() and not block.lstrip().startswith("#")
    ]


def _final_text_units(text: str, output_language: str) -> int:
    if output_language == "English":
        return len(re.findall(r"\b[\w]+(?:[-'][\w]+)*\b", text, re.UNICODE))
    han = len(re.findall(r"[\u3400-\u9fff]", text))
    without_han = re.sub(r"[\u3400-\u9fff]", " ", text)
    words = len(re.findall(r"\b[\w]+(?:[-'][\w]+)*\b", without_han, re.UNICODE))
    return han + words


def _body_false_attributions(
    body: BodyDraftPackage,
) -> list[tuple[str, int, str]]:
    prose = _body_prose_paragraphs(body.markdown)
    targets = list(
        dict.fromkeys(
            (citation.section_id, citation.paragraph_number)
            for citation in body.citations
        )
    )
    findings: list[tuple[str, int, str]] = []
    for (section_id, paragraph_number), paragraph in zip(targets, prose, strict=False):
        detail = false_self_attribution_detail(paragraph)
        if detail:
            findings.append((section_id, paragraph_number, detail))
    return findings


def _final_matter_error_detail(exc: ValidationError | FinalDeliveryError) -> str:
    if isinstance(exc, ValidationError):
        return str(exc.errors(include_url=False)[:8])
    return str(exc)


def _canonical_section_kind(value: str) -> str | None:
    normalized = re.sub(r"\s+", "", value.casefold())
    if "存在问题" in normalized and "技术发展趋势" in normalized:
        return "problems_and_trends"
    aliases = {
        "引言": "introduction",
        "introduction": "introduction",
        "国内外研究现状": "research_status",
        "研究现状": "research_status",
        "researchstatus": "research_status",
        "现状分析": "current_status_analysis",
        "currentstatusanalysis": "current_status_analysis",
        "存在问题": "problems",
        "existingproblems": "problems",
        "技术发展趋势": "trends",
        "technologytrends": "trends",
        "结语": "conclusion",
        "结论": "conclusion",
        "conclusion": "conclusion",
    }
    return aliases.get(normalized)


def _demote_body_headings(body: str) -> str:
    return re.sub(r"^##\s+", "### ", body, flags=re.MULTILINE)


def _font_size(value: str | None) -> float:
    if not value:
        return 11.0
    mapping = {"小四": 12.0, "四号": 14.0, "五号": 10.5, "小五": 9.0}
    if value.strip() in mapping:
        return mapping[value.strip()]
    match = re.search(r"(\d+(?:\.\d+)?)", value)
    return float(match.group(1)) if match else 11.0


def _set_run_font(run, font_name: str, size: float, *, bold: bool = False) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _add_section(document, title: str, content: str, *, level: int) -> None:
    document.add_heading(title, level=level)
    for paragraph in _split_prose_paragraphs(content):
        document.add_paragraph(paragraph)


def _append_markdown(
    document,
    markdown: str,
    font_name: str,
    font_size: float,
    *,
    numeric_superscript: bool = False,
) -> None:
    for block in re.split(r"\n\s*\n", markdown.strip()):
        clean = block.strip()
        if not clean:
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", clean)
        if heading:
            document.add_heading(heading.group(2).strip(), level=len(heading.group(1)))
            continue
        paragraph = document.add_paragraph()
        text = re.sub(r"\*\*([^*]+)\*\*", r"\1", clean)
        if not numeric_superscript:
            run = paragraph.add_run(text)
            _set_run_font(run, font_name, font_size)
            continue
        _append_numeric_citation_runs(paragraph, text, font_name, font_size)


def _append_numeric_citation_runs(
    paragraph,
    text: str,
    font_name: str,
    font_size: float,
) -> None:
    """Render code-owned numeric citation clusters as Word superscripts."""

    citation_pattern = re.compile(
        r"\[(?:\d+(?:,\s*[^;\]\[]+)?)(?:;\s*\d+(?:,\s*[^;\]\[]+)?)*\s*\]"
    )
    cursor = 0
    for match in citation_pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            _set_run_font(run, font_name, font_size)
        citation_run = paragraph.add_run(match.group(0))
        _set_run_font(citation_run, font_name, font_size)
        citation_run.font.superscript = True
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        _set_run_font(run, font_name, font_size)
