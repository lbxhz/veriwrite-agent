"""Independent DOCX/PDF paper extraction and quality evaluation."""

from __future__ import annotations

import hashlib
import json
import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from docx import Document
from pydantic import ValidationError
from pypdf import PdfReader

from veriwrite_agent.llm.base import LLMClient
from veriwrite_agent.models.standalone_evaluation import (
    StandalonePaperEvaluation,
    StandaloneQualityMetric,
    StandaloneSemanticReview,
)

MAX_UPLOAD_BYTES = 30 * 1024 * 1024
MAX_REVIEW_CHARACTERS = 70_000


@dataclass(frozen=True)
class ExtractedPaper:
    filename: str
    source_format: str
    text: str
    page_count: int | None
    extraction_method: str


@dataclass(frozen=True)
class CitationProfile:
    reference_count: int
    citation_marker_count: int
    matched_marker_count: int
    cited_reference_count: int
    doi_count: int
    score: float
    basis: tuple[str, ...]


def extract_uploaded_paper(filename: str, payload: bytes) -> ExtractedPaper:
    """Extract paper text without storing the uploaded source file."""

    if not payload:
        raise ValueError("上传文件为空")
    if len(payload) > MAX_UPLOAD_BYTES:
        raise ValueError("单个论文文件不能超过 30 MB")
    suffix = Path(filename).suffix.casefold()
    if suffix == ".docx":
        try:
            document = Document(BytesIO(payload))
        except Exception as exc:
            raise ValueError("DOCX 文件损坏或无法读取") from exc
        lines: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = str(getattr(paragraph.style, "name", "") or "")
            if style_name.casefold().startswith("heading"):
                text = f"## {text}"
            lines.append(text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))
        text = "\n\n".join(lines)
        method = "DOCX 原生文本"
        page_count = None
        source_format = "docx"
    elif suffix == ".pdf":
        try:
            reader = PdfReader(BytesIO(payload))
            pages = [(page.extract_text() or "").strip() for page in reader.pages]
        except Exception as exc:
            raise ValueError("PDF 文件损坏、加密或无法读取") from exc
        if not pages:
            raise ValueError("PDF 不包含可读取页面")
        empty_pages = [index + 1 for index, page in enumerate(pages) if not page]
        if len(empty_pages) == len(pages):
            raise ValueError("该 PDF 是扫描件且没有可提取文字；请先进行 OCR 后再评测")
        text = "\n\n".join(
            f"[第 {index} 页]\n{page}"
            for index, page in enumerate(pages, 1)
            if page
        )
        method = "PDF 分页原生文本"
        if empty_pages:
            method += f"（{len(empty_pages)} 页无文本，未纳入评测）"
        page_count = len(pages)
        source_format = "pdf"
    else:
        raise ValueError("仅支持 .docx 和 .pdf 论文")
    if len(re.sub(r"\s+", "", text)) < 300:
        raise ValueError("可提取正文不足 300 个字符，无法进行可靠评测")
    return ExtractedPaper(
        filename=Path(filename).name,
        source_format=source_format,
        text=text,
        page_count=page_count,
        extraction_method=method,
    )


class StandalonePaperEvaluationService:
    """Combine deterministic document checks with an independent LLM judge."""

    _WEIGHTS = {
        "requirement_compliance": 0.15,
        "citation_consistency": 0.25,
        "topic_focus": 0.15,
        "analysis_synthesis": 0.20,
        "structure_organization": 0.15,
        "language_style": 0.10,
    }
    _LABELS = {
        "requirement_compliance": "要求符合度",
        "citation_consistency": "引文编号与论断匹配",
        "topic_focus": "主题聚焦",
        "analysis_synthesis": "分析与综合",
        "structure_organization": "结构与组织",
        "language_style": "语言与学术表达",
    }

    def __init__(self, client: LLMClient, *, reviewer_model: str) -> None:
        self._client = client
        self._reviewer_model = reviewer_model

    def evaluate(
        self,
        paper: ExtractedPaper,
        *,
        expected_topic: str = "",
        requirements: str = "",
    ) -> StandalonePaperEvaluation:
        body, references = _split_references(paper.text)
        citation_profile = _citation_profile(body, references)
        review = self._semantic_review(
            body,
            references,
            expected_topic=expected_topic,
            requirements=requirements,
        )
        semantic_by_code = {metric.code: metric for metric in review.metrics}
        metrics: list[StandaloneQualityMetric] = []
        for code, weight in self._WEIGHTS.items():
            if code == "citation_consistency":
                semantic = semantic_by_code[code]
                score = 0.35 * citation_profile.score + 0.65 * semantic.score
                if any(
                    finding.dimension == "citation_consistency"
                    and finding.severity == "major"
                    for finding in review.findings
                ):
                    score = min(score, 35.0)
                basis = [
                    *citation_profile.basis,
                    "论断—文献语义抽查：" + semantic.rationale,
                ]
            else:
                semantic = semantic_by_code[code]
                score = semantic.score
                basis = [semantic.rationale]
            rounded = round(max(0.0, min(100.0, score)), 2)
            metrics.append(
                StandaloneQualityMetric(
                    code=code,
                    label=self._LABELS[code],
                    score=rounded,
                    weight=weight,
                    weighted_points=round(rounded * weight, 2),
                    basis=basis,
                )
            )
        overall = round(sum(metric.weighted_points for metric in metrics), 2)
        criteria = json.dumps(
            {
                "expected_topic": expected_topic.strip(),
                "requirements": requirements.strip(),
                "method": "document-quality-judge-v2",
            },
            ensure_ascii=False,
            sort_keys=True,
        )
        return StandalonePaperEvaluation(
            paper_fingerprint=hashlib.sha256(paper.text.encode("utf-8")).hexdigest(),
            criteria_fingerprint=hashlib.sha256(criteria.encode("utf-8")).hexdigest(),
            source_filename=paper.filename,
            source_format=paper.source_format,
            extraction_method=paper.extraction_method,
            page_count=paper.page_count,
            counted_units=_count_units(body),
            reference_count=citation_profile.reference_count,
            citation_marker_count=citation_profile.citation_marker_count,
            inferred_title=review.inferred_title,
            inferred_topic=review.inferred_topic,
            overall_score=overall,
            grade=_grade(overall),
            metrics=metrics,
            findings=review.findings,
            reviewer_model=self._reviewer_model,
            limitations=[
                "独立入口只能检查成品文档，无法验证作者是否真正阅读了来源。",
                "论断—文献匹配只依据正文语境与文后题名抽查；未同时提供来源 PDF 与证据库，"
                "因此不能代替逐句事实蕴含和来源页码核验。",
                "LLM 评分适合比较和定位问题，不等同于教师、期刊或查重系统的最终判定。",
                "公平比较应使用相同主题、相同要求、相同评测器版本。",
            ],
        )

    def _semantic_review(
        self,
        body: str,
        references: str,
        *,
        expected_topic: str,
        requirements: str,
    ) -> StandaloneSemanticReview:
        schema = json.dumps(
            StandaloneSemanticReview.model_json_schema(),
            ensure_ascii=False,
            separators=(",", ":"),
        )
        review_text = _review_sample(body)
        payload = {
            "expected_topic": expected_topic.strip() or "未指定；从论文中推断",
            "requirements": requirements.strip() or "按通用课程论文规范评价",
            "paper_text": review_text,
            "text_is_sampled": len(review_text) < len(body),
            "citation_contexts": _citation_contexts(body, references),
        }
        messages = [
            {
                "role": "system",
                "content": (
                    "你是独立的学术论文质量评审员。只评价给出的成品论文，不替作者补写，"
                    "也不要因为措辞可优化就虚构严重问题。分别评价：要求符合度、引文编号与论断匹配、主题聚焦、"
                    "分析与综合、结构与组织、语言与学术表达。要求符合度在没有专门要求时，"
                    "按通用课程论文的题目、摘要、关键词、问题意识、正文推进、结论等规范评价。"
                    "分析与综合重点判断论文是否围绕问题比较多项研究、归纳共识/差异/局限并形成"
                    "作者判断，而不是逐篇摘要串联。主题聚焦要识别真实跑题，不因合理背景材料扣分。"
                    "结构分关注章节职责、论证推进和重复；语言分关注自然、统一、无内部工作流指令。"
                    "引文匹配必须结合 citation_contexts 检查：编号存在不等于引用正确；若段落所述对象、"
                    "数据集、方法或年份与对应文后题名明显不一致，必须列为 citation_consistency 的 major finding。"
                    "这里只根据成品中可见信息做一致性审查，不臆测论文全文内容，也不核验文献是否真实。"
                    "每项给出0到100分和可核查理由；findings最多12项，只列影响明显的问题，"
                    "location必须指向章节、页码或可识别段落。引用编号覆盖和格式由代码计算，"
                    "你负责判断可见的论断—题名是否相符。输出自然中文 JSON，并严格满足模式："
                    f"{schema}"
                ),
            },
            {"role": "user", "content": json.dumps(payload, ensure_ascii=False)},
        ]
        last_error: Exception | None = None
        for attempt in range(3):
            raw = self._client.complete(
                messages,
                response_format={"type": "json_object"},
            )
            try:
                return StandaloneSemanticReview.model_validate_json(raw)
            except (ValidationError, ValueError) as exc:
                last_error = exc
                if attempt < 2:
                    messages.extend(
                        [
                            {"role": "assistant", "content": raw},
                            {
                                "role": "user",
                                "content": (
                                    "只修复 JSON 合同，不改变评价结论。必须恰好包含六个语义指标，"
                                    "其中必须包含 citation_consistency。错误："
                                    f"{str(exc)[:700]}"
                                ),
                            },
                        ]
                    )
        raise ValueError(f"独立论文评审输出不符合数据合同：{last_error}")


def _split_references(text: str) -> tuple[str, str]:
    pattern = re.compile(
        r"(?im)^\s*(?:#{1,4}\s*)?(?:参考文献|references|bibliography)\s*[:：]?\s*$"
    )
    matches = list(pattern.finditer(text))
    if not matches:
        return text, ""
    marker = matches[-1]
    return text[: marker.start()].strip(), text[marker.end() :].strip()


def _citation_contexts(body: str, references: str) -> list[dict[str, object]]:
    """Pair visible numeric citations with their surrounding claims and entries."""

    reference_entries = {
        int(match.group(1)): re.sub(r"\s+", " ", match.group(2)).strip()
        for match in re.finditer(
            r"(?m)^\s*\[(\d{1,4})\]\s*(.+?)\s*$",
            references,
        )
    }
    contexts: list[dict[str, object]] = []
    for paragraph_number, paragraph in enumerate(
        (item.strip() for item in re.split(r"\n\s*\n", body)),
        1,
    ):
        if not paragraph:
            continue
        numbers = list(
            dict.fromkeys(
                number
                for group in re.findall(r"\[([^\]\n]{1,300})\]", paragraph)
                for number in _citation_numbers(group)
            )
        )
        if not numbers:
            continue
        contexts.append(
            {
                "paragraph": paragraph_number,
                "claim_context": re.sub(r"\s+", " ", paragraph)[:1200],
                "cited_references": {
                    str(number): reference_entries.get(number, "<编号未映射>")
                    for number in numbers
                },
            }
        )
        if len(contexts) >= 80:
            break
    return contexts


def _citation_profile(body: str, references: str) -> CitationProfile:
    entries = re.findall(r"(?m)^\s*\[(\d{1,4})\]\s*.+$", references)
    reference_numbers = {int(number) for number in entries}
    if not entries and references:
        entries = re.findall(
            r"(?m)^\s*(?=.+(?:19|20)\d{2}).{20,}(?:\.|。)?\s*$",
            references,
        )
    reference_count = len(entries)
    doi_count = len(
        set(
            match.casefold().rstrip(".,;。；")
            for match in re.findall(r"10\.\d{4,9}/[-._;()/:A-Z0-9]+", references, re.I)
        )
    )
    bracket_groups = re.findall(r"\[([^\]\n]{1,300})\]", body)
    numeric_groups = [
        numbers
        for group in bracket_groups
        if (numbers := _citation_numbers(group))
    ]
    numeric_markers = [number for numbers in numeric_groups for number in numbers]
    author_year_markers = re.findall(
        r"[（(][^()（）\n]{1,60}(?:19|20)\d{2}[a-z]?[^()（）\n]{0,20}[）)]",
        body,
        re.I,
    )
    marker_count = len(numeric_groups) + len(author_year_markers)
    if reference_numbers and numeric_markers:
        matched = sum(number in reference_numbers for number in numeric_markers)
        cited = len(reference_numbers.intersection(numeric_markers))
        match_ratio = matched / len(numeric_markers)
        coverage = cited / len(reference_numbers)
        doi_ratio = min(doi_count / max(reference_count, 1), 1.0)
        score = 55 * match_ratio + 35 * coverage + 10 * doi_ratio
        basis = (
            f"数字引用标记可映射 {matched}/{len(numeric_markers)}",
            f"文后条目被正文引用 {cited}/{reference_count}",
            f"含 DOI 的参考文献 {doi_count}/{reference_count}",
        )
        matched_count = matched
        cited_count = cited
    elif reference_count and author_year_markers:
        density = min(len(author_year_markers) / reference_count, 1.0)
        doi_ratio = min(doi_count / reference_count, 1.0)
        score = 55 + 30 * density + 15 * doi_ratio
        basis = (
            f"检测到作者—年份引用 {len(author_year_markers)} 处",
            f"文后参考文献 {reference_count} 条",
            f"含 DOI 的参考文献 {doi_count}/{reference_count}",
        )
        matched_count = len(author_year_markers)
        cited_count = min(len(author_year_markers), reference_count)
    elif reference_count:
        score = 30 + 20 * min(doi_count / reference_count, 1.0)
        basis = (
            f"文后参考文献 {reference_count} 条，但未识别到正文引用标记",
            f"含 DOI 的参考文献 {doi_count}/{reference_count}",
        )
        matched_count = 0
        cited_count = 0
    else:
        score = 10.0
        basis = ("未识别到独立参考文献表", "未识别到可核对的正文引用标记")
        matched_count = 0
        cited_count = 0
    return CitationProfile(
        reference_count=reference_count,
        citation_marker_count=marker_count,
        matched_marker_count=matched_count,
        cited_reference_count=cited_count,
        doi_count=doi_count,
        score=score,
        basis=basis,
    )


def _citation_numbers(group: str) -> list[int]:
    """Read reference numbers while ignoring page locators inside one marker."""

    values: list[int] = []
    for segment in re.split(r"\s*[;；]\s*", group):
        match = re.match(r"\s*(\d{1,4})(?:\s*[-–]\s*(\d{1,4}))?", segment)
        if match is None:
            continue
        start = int(match.group(1))
        if start < 1:
            # Mathematical intervals such as [0,1] are not reference markers.
            continue
        end_text = match.group(2)
        if end_text is None:
            values.append(start)
            continue
        end = int(end_text)
        if start <= end <= start + 100:
            values.extend(range(start, end + 1))
    return values


def _review_sample(body: str) -> str:
    if len(body) <= MAX_REVIEW_CHARACTERS:
        return body
    head = body[:25_000]
    middle_start = max((len(body) - 20_000) // 2, 25_000)
    middle = body[middle_start : middle_start + 20_000]
    tail = body[-25_000:]
    return f"{head}\n\n[中部等距抽样]\n{middle}\n\n[末尾抽样]\n{tail}"


def _count_units(text: str) -> int:
    han = len(re.findall(r"[\u3400-\u9fff]", text))
    words = len(re.findall(r"\b[A-Za-z]+(?:[-'][A-Za-z]+)*\b", text))
    return han + words


def _grade(score: float) -> str:
    if score >= 90:
        return "excellent"
    if score >= 80:
        return "strong"
    if score >= 70:
        return "acceptable"
    return "weak"
