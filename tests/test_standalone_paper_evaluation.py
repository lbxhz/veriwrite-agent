import json
from io import BytesIO

import pytest
from docx import Document

from veriwrite_agent.services.standalone_paper_evaluation import (
    StandalonePaperEvaluationService,
    _citation_profile,
    extract_uploaded_paper,
)


class _FakeReviewClient:
    def complete(self, messages, *, response_format=None):
        assert response_format == {"type": "json_object"}
        payload = json.loads(messages[1]["content"])
        assert payload["citation_contexts"]
        assert payload["citation_contexts"][0]["cited_references"]["1"]
        return json.dumps(
            {
                "inferred_title": "大气遥感综述",
                "inferred_topic": "大气遥感的数据获取与反演",
                "metrics": [
                    {
                        "code": "requirement_compliance",
                        "score": 88,
                        "rationale": "包含完整课程论文结构。",
                    },
                    {
                        "code": "citation_consistency",
                        "score": 90,
                        "rationale": "正文论断与文后题名相符。",
                    },
                    {
                        "code": "topic_focus",
                        "score": 90,
                        "rationale": "正文围绕大气遥感展开。",
                    },
                    {
                        "code": "analysis_synthesis",
                        "score": 84,
                        "rationale": "能够比较研究差异并归纳局限。",
                    },
                    {
                        "code": "structure_organization",
                        "score": 86,
                        "rationale": "章节职责清楚且有递进。",
                    },
                    {
                        "code": "language_style",
                        "score": 92,
                        "rationale": "中文表达统一自然。",
                    },
                ],
                "findings": [],
            },
            ensure_ascii=False,
        )


def _docx_bytes() -> bytes:
    document = Document()
    document.add_heading("大气遥感综述", level=1)
    document.add_heading("摘要", level=2)
    document.add_paragraph("大气遥感是认识天气与气候过程的重要技术。" * 20)
    document.add_heading("引言", level=2)
    document.add_paragraph("不同观测平台具有互补优势，相关研究形成了多种技术路线[1]。" * 12)
    document.add_heading("参考文献", level=2)
    document.add_paragraph("[1] Zhang. Atmospheric remote sensing. 2024. DOI:10.1000/test")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_docx_can_be_evaluated_without_project_state() -> None:
    paper = extract_uploaded_paper("paper.docx", _docx_bytes())

    report = StandalonePaperEvaluationService(
        _FakeReviewClient(), reviewer_model="fake-reviewer"
    ).evaluate(paper, expected_topic="大气遥感")

    assert report.source_format == "docx"
    assert report.reference_count == 1
    assert report.citation_marker_count >= 1
    assert len(report.metrics) == 6
    assert report.overall_score > 80


def test_pdf_loader_accepts_machine_readable_pdf(monkeypatch) -> None:
    class _Page:
        def extract_text(self):
            return "论文标题\n摘要\n" + "可提取的论文正文内容。" * 40

    class _Reader:
        pages = [_Page(), _Page()]

    monkeypatch.setattr(
        "veriwrite_agent.services.standalone_paper_evaluation.PdfReader",
        lambda _: _Reader(),
    )

    paper = extract_uploaded_paper("paper.pdf", b"%PDF-1.7\nnot-empty")

    assert paper.source_format == "pdf"
    assert paper.page_count == 2
    assert "[第 1 页]" in paper.text


def test_scanned_pdf_without_text_gives_actionable_error(monkeypatch) -> None:
    class _Page:
        def extract_text(self):
            return ""

    class _Reader:
        pages = [_Page()]

    monkeypatch.setattr(
        "veriwrite_agent.services.standalone_paper_evaluation.PdfReader",
        lambda _: _Reader(),
    )

    with pytest.raises(ValueError, match="扫描件"):
        extract_uploaded_paper("scan.pdf", b"%PDF-1.7\nnot-empty")


def test_invalid_extension_is_rejected() -> None:
    with pytest.raises(ValueError, match="仅支持"):
        extract_uploaded_paper("paper.txt", b"content")


def test_page_locators_do_not_hide_numeric_citations() -> None:
    profile = _citation_profile(
        "研究形成了共识[1, pp. 2, 3; 2; 3, p. 8]。",
        "\n".join(
            [
                "[1] Alpha. Study. 2022. DOI:10.1000/a",
                "[2] Beta. Study. 2023. DOI:10.1000/b",
                "[3] Gamma. Study. 2024. DOI:10.1000/c",
            ]
        ),
    )

    assert profile.citation_marker_count == 1
    assert profile.matched_marker_count == 3
    assert profile.cited_reference_count == 3
    assert profile.score == 100


def test_mathematical_zero_to_one_interval_is_not_a_citation() -> None:
    profile = _citation_profile(
        "模型输入归一化到[0,1]，相关方法见[1]。",
        "[1] Alpha. Study. 2022. DOI:10.1000/a",
    )

    assert profile.citation_marker_count == 1
    assert profile.matched_marker_count == 1
    assert profile.cited_reference_count == 1
